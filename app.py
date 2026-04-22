import streamlit as st
from docx import Document
import PyPDF2
import pdfplumber
import pyperclip
import json
import re
import io

# Character mappings
mapping_standard = {
    "`": "ä",
    "1": "a̱",
    "2": "ë",
    "3": "e̱",
    "5": "i̱",
    "6": "ö",
    "7": "o̱",
    "0": "ɛ̈",
    "]": "ɔ̱",
    "s": "ɔ",
    "f": "ɣ",
    "x": "ŋ",
    "v": "ɛ",

    # Uppercase mappings
    "~": "Ä",
    "!": "A̱",
    "@": "Ë",
    "#": "E̱",
    "%": "I̱",
    "^": "Ö",
    "&": "O̱",
    ")": "Ɛ̈",
    "}": "Ɔ̱",
    "S": "Ɔ",
    "F": "Ɣ",
    "X": "Ŋ",
    "V": "Ɛ"
}

mapping_bok_en_yel = {
    "Æ": "ɛ",
    "Å": "ɔ̱",
    "À": "a̱",
    "Ÿ": "ɛ̈",
    "È": "e̱",
    "Œ": "ɔ",
    "Ŕ": "ɣ",
    "Ì": "i̱",
    "Ñ": "ŋ",
    "Ò": "o̱",
    "Ý": "ɛ̱̈",
    # Lowercase versions
    "æ": "ɛ",
    "å": "ɔ̱",
    "à": "a̱",
    "ÿ": "ɛ̈",
    "è": "e̱",
    "œ": "ɔ",
    "ŕ": "ɣ",
    "ì": "i̱",
    "ñ": "ŋ",
    "ò": "o̱",
    "ý": "ɛ̱̈",
}

def convert_text(text, char_mapping):
    protected = {}
    counter = 0

    # ---------- Protect Bible references ----------
    pattern = r"\b\d+\s*:\s*\d+[a-zA-Z]?(?:\s*-\s*\d+[a-zA-Z]?)*(?:\s*,\s*\d+[a-zA-Z]?)*\b"
    matches = list(re.finditer(pattern, text))
    for match in reversed(matches):
        key = f"__REF{chr(65 + counter)}__"
        protected[key] = match.group()
        start, end = match.span()
        text = text[:start] + key + text[end:]
        counter += 1
    
    # Protect decimals
    pattern_decimal = r"\b\d+\.\d+\b"
    matches = list(re.finditer(pattern_decimal, text))

    for match in reversed(matches):
        key = f"__DEC{counter}__"
        protected[key] = match.group()
        start, end = match.span()
        text = text[:start] + key + text[end:]
        counter += 1

    # ---------- Protect numbers with 4+ digits ----------
    pattern_big_numbers = r"\b\d{4,}\b"
    matches = list(re.finditer(pattern_big_numbers, text))
    for match in reversed(matches):
        key = f"__NUM{chr(65 + counter)}__"
        protected[key] = match.group()
        start, end = match.span()
        text = text[:start] + key + text[end:]
        counter += 1

    # ---------- Convert characters ----------
    result = ""
    i = 0
    open_paren = 0
    open_brack = 0
    punctuation = '.,;:—!?"  \n'
    lowercase_like = "abcdefghijklmnopqrstuvwxyz`1235670]sfxv" 
    
    while i < len(text):

        # Skip ANY placeholder (__REFX__ or __NUMX__)
        if text[i:i+2] == "__":
            end_index = text.find("__", i+2) + 2
            result += text[i:end_index]
            i = end_index
            continue

        char = text[i]
        # ---------- Preserve newlines ----------
        if char == "\n":
            result += "\n"
            i += 1
            continue

        if char == "!":
            prev_char = text[i-1] if i-1 >= 0 else ""
            prev_result_char = result[-1] if result else ""

            # Treat as real "!" if previous char is lowercase-like OR previous result char is a preserved "!"
            if prev_char in lowercase_like or prev_result_char == "!":
                result += "!"
            else:
                # Otherwise treat as fake letter
                result += char_mapping.get(char, char)

            i += 1
            continue

        # ---------- Track open parentheses ----------
        if char == "(":
            open_paren += 1
            result += char
            i += 1
            continue

        if char == ")" and open_paren > 0:
            prev_char = text[i-1] if i-1 >= 0 else ""
            next_char = text[i+1] if i+1 < len(text) else ""

            if ((prev_char in punctuation or next_char in punctuation) and
                not (prev_char.isalpha() and next_char.isalpha())):
                open_paren -= 1
                result += char
                i += 1
                continue

        # ---------- Track open brackets ----------
        if char == "[":
            open_brack += 1
            result += char
            i += 1
            continue

        if char == "]" and open_brack > 0:
            prev_char = text[i-1] if i-1 >= 0 else ""
            next_char = text[i+1] if i+1 < len(text) else ""

            if ((prev_char in punctuation or next_char in punctuation) and
                not (prev_char.isalpha() and next_char.isalpha())):
                open_brack -= 1
                result += char
                i += 1
                continue

        # ---------- Repeated fake vowels ----------
        if i + 1 < len(text) and text[i] == text[i + 1] and text[i] in char_mapping:
            result += char_mapping[text[i]] + char_mapping[text[i + 1]]
            i += 2
            continue

        # ---------- Normal conversion ----------
        result += char_mapping.get(char, char)
        i += 1

    # ---------- Restore protected parts ----------
    for key, value in protected.items():
        result = result.replace(key, value)

    return result

# -------- FILE READERS --------

def read_txt(file):
    return file.read().decode("utf-8")

def read_docx(file):
    doc = Document(file)
    text = ""
    for p in doc.paragraphs:
        text += p.text + "\n"
    return text

def read_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


# -------- FILE EXPORTERS --------

def export_txt(text):
    return text.encode("utf-8")

def export_docx(text):
    buffer = io.BytesIO()
    doc = Document()

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(buffer)
    buffer.seek(0)

    return buffer


# -------- STREAMLIT UI --------

st.markdown(
    "<h1 style='text-align: center;'>Nuer Font Converter</h1>",
    unsafe_allow_html=True
)

st.markdown(
    "<p style='text-align: center;'>Convert fake Nuer fonts to standard Unicode Nuer letters</p>",
    unsafe_allow_html=True
)

st.markdown("---")  # horizontal line

st.info(
    "Converting texts from fake Nuer fonts to **standard Unicode Nuer letters** is important for:\n\n"
    "- Making your texts **digitally searchable**.\n"
    "- Ensuring **compatibility across devices and platforms**.\n"
    "- Preserving **accurate letters**.\n\n"
    "This converter automatically handles common fake font mappings, preserves Bible references, "
    "numbers, and brackets. Please review your text manually if any numbers or words need adjustment."
)

st.markdown("---")  # horizontal line

# Font type selector
font_type = st.radio(
    "Select the fake font type used in your text",
    ["Standard Nuer Font", "Bok Ɛn Yel Font"],
    horizontal=True
)

if font_type == "Standard Nuer Font":
    active_mapping = mapping_standard
else:
    active_mapping = mapping_bok_en_yel

# Paste text
text_input = st.text_area("Paste fake-font Nuer text")

# Upload file
uploaded_file = st.file_uploader(
    "Or upload a file",
    type=["txt", "docx", "pdf"]
)

text = ""

if text_input:
    text = text_input

elif uploaded_file:

    if uploaded_file.name.endswith(".txt"):
        text = read_txt(uploaded_file)

    elif uploaded_file.name.endswith(".docx"):
        text = read_docx(uploaded_file)

    elif uploaded_file.name.endswith(".pdf"):
        text = read_pdf(uploaded_file)

# ---------- Generate Button ----------
if st.button("Generate Converted Text") and text:

    converted = convert_text(text, active_mapping)
    st.session_state.converted = converted  # save converted text in session

    st.subheader("Converted Text")
    st.text_area("Result", converted, height=300)

# ---------- Copy to Clipboard Button ----------
import streamlit.components.v1 as components

# ---------- Copy to Clipboard ----------
if 'converted' in st.session_state and st.session_state.converted:
    converted_text = st.session_state.converted

    components.html(f"""
    <button onclick="copyText()">Copy to Clipboard</button>

    <script>
    function copyText() {{
        const text = `{converted_text}`;
        navigator.clipboard.writeText(text);
        alert("Copied to clipboard!");
    }}
    </script>
    """, height=50)

# ---------- Always-visible download ----------
st.subheader("Download Converted File")
format_choice = st.selectbox("Choose download format", ["DOCX", "TXT"])

# Determine base filename
if uploaded_file:
    base_name = uploaded_file.name.rsplit(".", 1)[0]  # remove extension
elif text_input:
    base_name = "converted_text"
else:
    base_name = "converted_nuer"

# Enable download only if conversion is done
if 'converted' in st.session_state and st.session_state.converted:
    converted_text = st.session_state.converted

    if format_choice == "TXT":
        file_data = export_txt(converted_text)
        file_name = f"{base_name}_converted.txt"
    else:
        file_data = export_docx(converted_text)
        file_name = f"{base_name}_converted.docx"

    st.download_button(
        "Download Converted File",
        data=file_data,
        file_name=file_name
    )
else:
    st.download_button(
        "Download Converted File",
        data=b"",
        file_name=f"{base_name}_converted.txt",
        disabled=True
    )

st.warning(
    "After conversion, please manually review the text for the following cases:\n\n"
    "- Words that were originally meant to remain in **English**.\n"
    "- **Single numbers between 0 and 999**, which may have been used as fake-font letters.\n"
    "- Cases where **ɛ̈ was manually underlined in MS Word** (not Unicode underline) to represent **ɛ̱̈**.\n\n"
    "These situations cannot always be detected automatically and may need manual correction."
)
# ---------- Footer / Notes ----------

st.markdown("---")  # horizontal line

st.markdown("### Note")
st.markdown("This converter is specifically designed for Nuer texts written using **fake fonts**.")

if font_type == "Standard Nuer Font":
    st.markdown(
        "**Standard Nuer Font mapping:**\n\n"
        "- `` ` `` → ä\n"
        "- `` 1 `` → a̱\n"
        "- `` 2 `` → ë\n"
        "- `` 3 `` → e̱\n"
        "- `` 5 `` → i̱\n"
        "- `` 6 `` → ö\n"
        "- `` 7 `` → o̱\n"
        "- `` 0 `` → ɛ̈\n"
        "- `` ] `` → ɔ̱\n"
        "- `` s `` → ɔ\n"
        "- `` f `` → ɣ\n"
        "- `` x `` → ŋ\n"
        "- `` v `` → ɛ\n\n"
    )
else:
    st.markdown(
        "**Bok Ɛn Yel Font mapping:**\n\n"
        "- `` Æ/æ `` → ɛ\n"
        "- `` Å/å `` → ɔ̱\n"
        "- `` À/à `` → a̱\n"
        "- `` Ÿ/ÿ `` → ɛ̈\n"
        "- `` È/è `` → e̱\n"
        "- `` Œ/œ `` → ɔ\n"
        "- `` Ŕ/ŕ `` → ɣ\n"
        "- `` Ì/ì `` → i̱\n"
        "- `` Ñ/ñ `` → ŋ\n"
        "- `` Ò/ò `` → o̱\n"
        "- `` Ý/ý `` → ɛ̱̈\n\n"
    )
st.markdown("Please note:\n"
    "The converter may not work correctly for other fonts or encodings.\n"
)

st.markdown("---")  # horizontal line
st.markdown(
    "### Credits\n"
    "Developed by **Jack Bill Jack**. Inspired by the need to preserve and digitize Nuer texts accurately.\n" 
    )
st.markdown(
    "Contact: jackbilljack14@gmail.com"
    )
st.markdown(
    "Special thanks to the Nuer community for font documentation and testing."
)

st.markdown("---")  # horizontal line
st.markdown(
    "For any suggestions or comments, please reach out to me through the email address above.\n"

    )
